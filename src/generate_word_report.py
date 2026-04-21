from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path
from typing import Dict, List

import cv2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def parse_summary_table(summary_path: Path) -> Dict[str, str]:
    lines = [line.rstrip("\n") for line in summary_path.read_text(encoding="utf-8", errors="ignore").splitlines() if line.strip()]
    if len(lines) < 2:
        return {}

    header = lines[0].split()

    combined_line = ""
    for line in reversed(lines):
        if line.strip().startswith("COMBINED"):
            combined_line = line
            break

    metrics: Dict[str, str] = {}

    if combined_line:
        values = combined_line.split()
        if values and values[0] == "COMBINED":
            for index, key in enumerate(header[1:], start=1):
                if index < len(values):
                    metrics[key] = values[index]
            if metrics:
                return metrics

    # 兼容 TrackEval 常见两行格式：第一行为字段名，第二行为数值（无 COMBINED 列）
    values = lines[1].split()
    for index, key in enumerate(header):
        if index < len(values):
            metrics[key] = values[index]
    return metrics


def load_trackeval_metrics(trackeval_dir: Path) -> Dict[str, str]:
    metric_files = [
        trackeval_dir / "pedestrian_summary.txt",
        trackeval_dir / "summary.txt",
    ]

    if trackeval_dir.exists() and trackeval_dir.is_dir():
        metric_files.extend(trackeval_dir.rglob("pedestrian_summary.txt"))
        metric_files.extend(trackeval_dir.rglob("summary.txt"))

    merged: Dict[str, str] = {}
    visited = set()
    for file in metric_files:
        file = file.resolve()
        if file in visited:
            continue
        visited.add(file)
        if file.exists():
            merged.update(parse_summary_table(file))
    return merged


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="生成 Word 格式实验报告")
    parser.add_argument("--run-meta", required=True, help="run_meta.json 路径")
    parser.add_argument("--trackeval-dir", required=True, help="TrackEval 输出目录（包含 summary 文件）")
    parser.add_argument("--video-path", required=True, help="原始视频路径")
    parser.add_argument("--tracked-video", required=True, help="输出跟踪视频路径")
    parser.add_argument("--mot-result", required=True, help="MOT结果文件路径")
    parser.add_argument("--figures-dir", default="report/figures/exp2", help="抽帧图片保存目录")
    parser.add_argument("--num-frames", type=int, default=4, help="抽帧数量")
    parser.add_argument("--output", default="report/OPEN_APPLICATION_REPORT.docx", help="输出 docx 路径")
    return parser.parse_args()


def add_kv_table(document: Document, rows: List[List[str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "项目"
    header_cells[1].text = "内容"
    for key, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)


def extract_frames(video_path: Path, out_dir: Path, num_frames: int) -> List[Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    capture = cv2.VideoCapture(str(video_path))
    if not capture.isOpened():
        return []

    total_frames = int(capture.get(cv2.CAP_PROP_FRAME_COUNT))
    if total_frames <= 0:
        capture.release()
        return []

    valid_count = max(1, num_frames)
    indices = sorted({int((total_frames - 1) * i / max(1, valid_count - 1)) for i in range(valid_count)})
    saved: List[Path] = []

    for index, frame_id in enumerate(indices, start=1):
        capture.set(cv2.CAP_PROP_POS_FRAMES, frame_id)
        ok, frame = capture.read()
        if not ok:
            continue
        file_path = out_dir / f"frame_{index:02d}_id_{frame_id:05d}.jpg"
        cv2.imwrite(str(file_path), frame)
        saved.append(file_path)

    capture.release()
    return saved


def add_bullet_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.style = "List Bullet"


def add_metric_meaning_section(document: Document) -> None:
    document.add_heading("7.1 评估参数含义详解", level=2)
    document.add_paragraph("以下参数来自 TrackEval 的 HOTA/CLEAR/Identity 三类指标，用于从“检测正确性、关联稳定性、身份一致性”三个维度评估跟踪质量。")

    add_kv_table(
        document,
        [
            ["HOTA", "高阶跟踪精度，综合检测与关联能力；值越高表示整体跟踪质量越好。"],
            ["DetA", "检测准确度，衡量检测框与真值匹配质量。"],
            ["AssA", "关联准确度，衡量跨帧 ID 关联是否正确稳定。"],
            ["DetRe / DetPr", "检测召回率 / 检测精确率，反映漏检与误检情况。"],
            ["AssRe / AssPr", "关联召回率 / 关联精确率，反映轨迹连接完整性与纯净度。"],
            ["LocA", "定位准确度，反映匹配目标的空间定位精度（通常与 IoU 相关）。"],
            ["OWTA", "强调关联质量的综合指标，用于补充 HOTA 在关联维度上的解释。"],
            ["MOTA", "多目标跟踪准确率，综合 FP、FN、IDSW；值越高越好。"],
            ["MOTP", "多目标跟踪定位精度，关注匹配成功样本的位置误差。"],
            ["CLR_TP / CLR_FN / CLR_FP", "真正例 / 漏检 / 误检计数，是 CLEAR 指标的底层统计量。"],
            ["IDF1", "身份 F1 分数，综合 IDP 与 IDR；强调身份一致性。"],
            ["IDP / IDR", "身份精确率 / 身份召回率，衡量身份预测准确与覆盖。"],
            ["IDSW", "ID 切换次数；越小越好，表示轨迹身份更稳定。"],
            ["MT / PT / ML", "大部分跟踪 / 部分跟踪 / 大部分丢失的真值轨迹数量。"],
            ["Frag", "轨迹碎片数；越小越好，表示轨迹连续性更好。"],
            ["Dets / GT_Dets", "预测检测总数 / 真值检测总数。"],
            ["IDs / GT_IDs", "预测轨迹 ID 数 / 真值轨迹 ID 数。"],
        ],
    )

    document.add_paragraph("解释建议：")
    add_bullet_paragraph(document, "若 HOTA 高但 IDF1 低，通常表示检测好但跨帧身份维护不足。")
    add_bullet_paragraph(document, "若 MOTA 低且 CLR_FP 高，通常表示误检偏多，应提高置信阈值或优化检测器。")
    add_bullet_paragraph(document, "若 IDSW 与 Frag 偏高，通常表示遮挡场景下关联不稳定，可引入更强 ReID 特征。")


def main() -> None:
    args = parse_args()

    run_meta_path = Path(args.run_meta)
    trackeval_dir = Path(args.trackeval_dir)
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)

    run_meta = json.loads(run_meta_path.read_text(encoding="utf-8"))
    metrics = load_trackeval_metrics(trackeval_dir)
    tracked_video_path = Path(args.tracked_video)
    figures = extract_frames(tracked_video_path, Path(args.figures_dir), args.num_frames)

    document = Document()
    title = document.add_heading("车辆轨迹检测实验报告（开放与应用）", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f"日期：{date.today().isoformat()}")

    document.add_heading("1. 实验目标", level=1)
    document.add_paragraph("基于真实道路车辆视频，完成目标检测与多目标轨迹跟踪，输出可视化结果并进行标准化评测。")

    document.add_heading("2. 项目框架设计", level=1)
    document.add_paragraph("本项目采用“检测-跟踪-可视化-评测-报告”五阶段流水线，模块间通过标准文件接口解耦，便于替换算法组件和复现实验。")
    add_bullet_paragraph(document, "数据输入层：读取视频流并逐帧解码，输入为 `runs/exp2/test_1.mp4`。")
    add_bullet_paragraph(document, "检测层：`ultralytics` YOLO 模型输出每帧边界框、类别和置信度。")
    add_bullet_paragraph(document, "跟踪层：`ByteTrack` 将检测框关联为时序轨迹并维护稳定 ID。")
    add_bullet_paragraph(document, "可视化层：`supervision` 绘制边框、轨迹线和 ID 标签，输出 `tracked.mp4`。")
    add_bullet_paragraph(document, "评测层：转换为 MOTChallenge 结构后由 `TrackEval` 计算 HOTA/MOTA/IDF1。")
    add_bullet_paragraph(document, "报告层：自动汇总运行元数据与评测指标，生成 Word 报告用于提交。")

    document.add_heading("3. 各部分实现细节", level=1)
    document.add_paragraph("核心脚本及职责如下：")
    add_kv_table(
        document,
        [
            ["`src/tracking_pipeline.py`", "视频逐帧推理；检测结果输入 ByteTrack；导出跟踪视频、MOT 文本和运行元数据。"],
            ["`src/utils/mot_io.py`", "MOTChallenge 行格式定义与读写，统一跟踪结果持久化接口。"],
            ["`src/prepare_trackeval_demo.py`", "将结果文件映射到 TrackEval 目录结构，并生成 `seqinfo.ini` 与 `seqmaps`。"],
            ["`src/evaluate_trackeval.py`", "封装官方评测脚本参数，执行 HOTA/CLEAR/Identity 指标计算。"],
            ["`src/generate_word_report.py`", "汇总元数据、评测指标并插入结果图，输出可提交报告。"],
        ],
    )

    document.add_paragraph("关键实现点：")
    add_bullet_paragraph(document, "检测与跟踪解耦：任何能输出边界框的检测器都可替换 YOLO 接入 ByteTrack。")
    add_bullet_paragraph(document, "结果标准化：统一导出 MOT 格式，便于与学术评测工具对接。")
    add_bullet_paragraph(document, "可复现实验：运行参数（模型、阈值、设备、吞吐）固化在 `run_meta.json`。")
    add_bullet_paragraph(document, "自动报告：将指标与可视化样例联动输出，减少人工整理工作量。")

    document.add_heading("4. 实验环境与方法", level=1)
    add_kv_table(
        document,
        [
            ["算法方案", "YOLO + ByteTrack + TrackEval + supervision"],
            ["输入视频", args.video_path],
            ["输出视频", args.tracked_video],
            ["MOT结果", args.mot_result],
            ["检测模型", str(run_meta.get("model", ""))],
            ["运行设备", str(run_meta.get("device", ""))],
            ["置信度阈值", str(run_meta.get("conf", ""))],
        ],
    )

    document.add_heading("5. 运行结果（exp2）", level=1)
    add_kv_table(
        document,
        [
            ["处理帧数", str(run_meta.get("frames", ""))],
            ["输入FPS", str(run_meta.get("fps_input", ""))],
            ["总耗时（秒）", str(run_meta.get("elapsed_sec", ""))],
            ["处理吞吐FPS", str(run_meta.get("throughput_fps", ""))],
            ["平均每帧跟踪数", str(run_meta.get("avg_tracks_per_frame", ""))],
        ],
    )

    document.add_heading("6. exp2 结果抽帧示例", level=1)
    if figures:
        document.add_paragraph("以下图片由 `runs/exp2_rerun/tracked.mp4` 自动抽帧生成，用于展示车辆检测与轨迹跟踪效果：")
        for index, figure in enumerate(figures, start=1):
            document.add_paragraph(f"图 {index}：{figure.name}")
            document.add_picture(str(figure), width=Inches(6.2))
    else:
        document.add_paragraph("未能从跟踪视频抽帧，请检查 `--tracked-video` 路径是否正确。")

    document.add_heading("7. TrackEval 指标", level=1)
    if metrics:
        selected_keys = ["HOTA", "DetA", "AssA", "MOTA", "IDF1", "CLR_TP", "IDTP"]
        rows: List[List[str]] = []
        for key in selected_keys:
            if key in metrics:
                rows.append([key, metrics[key]])
        if not rows:
            rows = [[key, value] for key, value in metrics.items()]
        add_kv_table(document, rows)
        add_metric_meaning_section(document)
    else:
        document.add_paragraph("未检测到 summary 指标文件，请先运行 TrackEval。")

    document.add_heading("8. 结论与应用建议", level=1)
    document.add_paragraph(
        "本实验基于 exp2 视频完成了从推理、跟踪、评测到报告输出的闭环流程，"
        "验证了项目框架可复用、可扩展、可复现。"
    )
    add_bullet_paragraph(document, "工程应用建议：接入多路摄像头并扩展统计模块（车流量、车速、拥堵指数）。")
    add_bullet_paragraph(document, "算法优化建议：按场景微调检测器阈值、替换更强 ReID 以降低遮挡场景 ID 切换。")
    add_bullet_paragraph(document, "评测建议：引入人工标注 GT，避免“结果自生成 GT”导致指标上限偏高。")

    document.save(str(output))
    print(f"word-report-ready: {output}")


if __name__ == "__main__":
    main()
